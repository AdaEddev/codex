"""Tkinter-based Azure OpenAI coding assistant for interview transcripts.

The app highlights relevant excerpts in a `.docx` transcript using the
following categories:

A. Background & Context — Course structure and participant role including assessment format and delivery.
B. Feasibility & Practical Implementation — Practical, logistical, and administrative aspects of implementing oral assessment.
C. Validity & Learning Assurance — Evidence that the oral assessment measured intended learning outcomes.
D. Disciplinary Relevance — Fit between oral assessment and disciplinary norms, skills, and values.
E. Student Engagement & Observations — Student reactions, fairness, and inclusivity observations.
F. Reflection & Improvement — What worked, what did not, and what to change next time.
G. Sustainability & Future Use — Whether this approach can be maintained, scaled, or used long term.
H. Additional Insights — Open reflections, emergent, or unanticipated themes.

The script requires the following environment variables so it can communicate
with an Azure OpenAI deployment:

```
AZURE_OPENAI_API_KEY
AZURE_OPENAI_ENDPOINT
AZURE_OPENAI_DEPLOYMENT
AZURE_OPENAI_API_VERSION (optional, defaults to 2024-02-15-preview)
```

It keeps the original transcript text intact while applying color-coded
highlights to the passages identified by the model.
"""

from __future__ import annotations

import json
import os
import threading
from dataclasses import dataclass
from tkinter import BOTH, END, LEFT, RIGHT, StringVar, Tk, Button, Frame, Label, Text, filedialog, messagebox, ttk

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
from openai import AzureOpenAI
from openai import OpenAIError


load_dotenv()


# Category metadata ---------------------------------------------------------
CATEGORY_DETAILS = {
    "A": {
        "title": "Background & Context",
        "color": "FFF2CC",  # pastel yellow
    },
    "B": {
        "title": "Feasibility & Practical Implementation",
        "color": "DAEEF3",  # pale aqua
    },
    "C": {
        "title": "Validity & Learning Assurance",
        "color": "E2F0D9",  # mint
    },
    "D": {
        "title": "Disciplinary Relevance",
        "color": "FCE4D6",  # blush peach
    },
    "E": {
        "title": "Student Engagement & Observations",
        "color": "E4DFEC",  # light lavender
    },
    "F": {
        "title": "Reflection & Improvement",
        "color": "D9E1F2",  # periwinkle
    },
    "G": {
        "title": "Sustainability & Future Use",
        "color": "F2F2F2",  # soft gray
    },
    "H": {
        "title": "Additional Insights",
        "color": "D5E8D4",  # pastel green
    },
}


SYSTEM_PROMPT = """You are an analyst that codes interview transcripts into the
specified categories. Return JSON following this schema:

{
  "matches": [
    {
      "category": "A",
      "quotes": ["verbatim excerpt"]
    }
  ]
}

Only output text that exists verbatim in the transcript. DO NOT paraphrase or
rewrite any text. Use the categories:

A. Background & Context – Course structure and participant role including assessment format and delivery.
B. Feasibility & Practical Implementation – Practical, logistical, and administrative aspects of implementing oral assessment.
C. Validity & Learning Assurance – Evidence that the oral assessment measured intended learning outcomes.
D. Disciplinary Relevance – Fit between oral assessment and disciplinary norms, skills, and values.
E. Student Engagement & Observations – Student reactions, fairness, and inclusivity observations.
F. Reflection & Improvement – What worked, what did not, and what to change next time.
G. Sustainability & Future Use – Whether this approach can be maintained, scaled, or used long term.
H. Additional Insights – Open reflections, emergent, or unanticipated themes.
"""


@dataclass
class QuoteMatch:
    category: str
    quote: str


class TranscriptCoderApp:
    def __init__(self, root: Tk) -> None:
        self.root = root
        self.root.title("Azure Oral Assessment Coder")

        self.client = self._build_client()

        self.input_path = StringVar()
        self.output_path = StringVar()

        self._build_ui()

    # ------------------------------------------------------------------ UI --
    def _build_ui(self) -> None:
        padding = {"padx": 10, "pady": 5}

        # File selection frame
        file_frame = Frame(self.root)
        file_frame.pack(fill=BOTH, **padding)

        Label(file_frame, text="Input transcript (.docx)").pack(anchor="w")
        input_row = Frame(file_frame)
        input_row.pack(fill=BOTH)
        Label(input_row, textvariable=self.input_path, width=60, anchor="w").pack(side=LEFT, fill=BOTH, expand=True)
        Button(input_row, text="Browse", command=self._select_input).pack(side=RIGHT)

        Label(file_frame, text="Output file (.docx)").pack(anchor="w", pady=(10, 0))
        output_row = Frame(file_frame)
        output_row.pack(fill=BOTH)
        Label(output_row, textvariable=self.output_path, width=60, anchor="w").pack(side=LEFT, fill=BOTH, expand=True)
        Button(output_row, text="Save As", command=self._select_output).pack(side=RIGHT)

        # Legend
        legend_frame = Frame(self.root)
        legend_frame.pack(fill=BOTH, **padding)
        Label(legend_frame, text="Coding legend:").pack(anchor="w")
        for code, meta in CATEGORY_DETAILS.items():
            ttk.Label(legend_frame, text=f"{code}. {meta['title']}").pack(anchor="w")

        # Action buttons
        action_frame = Frame(self.root)
        action_frame.pack(fill=BOTH, **padding)
        self.process_button = Button(action_frame, text="Process Transcript", command=self._start_processing)
        self.process_button.pack(side=LEFT)

        # Log window
        self.log = Text(self.root, height=12)
        self.log.pack(fill=BOTH, expand=True, padx=10, pady=10)

    # ------------------------------------------------------------- Handlers --
    def _select_input(self) -> None:
        filename = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
        if filename:
            self.input_path.set(filename)

    def _select_output(self) -> None:
        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
        if filename:
            self.output_path.set(filename)

    def _start_processing(self) -> None:
        if not self.input_path.get():
            messagebox.showerror("Missing input", "Please choose a transcript file.")
            return
        if not self.output_path.get():
            messagebox.showerror("Missing output", "Please choose where to save the highlighted document.")
            return
        self.process_button.config(state="disabled")
        threading.Thread(target=self._process_document, daemon=True).start()

    # --------------------------------------------------------- Core logic --
    def _process_document(self) -> None:
        try:
            self._log("Loading transcript...")
            document = Document(self.input_path.get())
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

            if not paragraphs:
                raise ValueError("No text found in the document.")

            matches: list[QuoteMatch] = []
            for chunk in chunk_text(paragraphs):
                matches.extend(self._code_chunk(chunk))

            if not matches:
                self._log("No excerpts were returned by the model. Saving original document.")
            else:
                applied = self._apply_highlights(document, matches)
                self._log(f"Applied {applied} highlight(s).")

            self._append_legend(document)

            document.save(self.output_path.get())
            self._log(f"Saved coded transcript to {self.output_path.get()}.")
            messagebox.showinfo("Done", "Processing complete!")
        except Exception as exc:  # noqa: BLE001
            messagebox.showerror("Processing failed", str(exc))
            self._log(f"ERROR: {exc}")
        finally:
            self.process_button.config(state="normal")

    def _code_chunk(self, chunk: str) -> list[QuoteMatch]:
        self._log(f"Coding chunk with {len(chunk)} characters...")
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    "Identify exact quotations from this transcript and map them "
                    "to the categories. Only include verbatim matches. Transcript:\n\n"
                    + chunk
                ),
            },
        ]
        try:
            response = self.client.chat.completions.create(
                model=os.environ["AZURE_OPENAI_DEPLOYMENT"],
                temperature=0,
                response_format={"type": "json_object"},
                messages=messages,
            )
        except OpenAIError as exc:
            raise RuntimeError(
                "Azure OpenAI call failed. Please confirm your endpoint, "
                "deployment name, API version, and key."
            ) from exc
        content = response.choices[0].message.content
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as exc:  # noqa: BLE001
            raise ValueError(f"Model response was not valid JSON: {content}") from exc

        matches = []
        for item in parsed.get("matches", []):
            category = item.get("category", "").strip().upper()
            if category not in CATEGORY_DETAILS:
                continue
            for quote in item.get("quotes", []):
                quote = (quote or "").strip()
                if quote:
                    matches.append(QuoteMatch(category=category, quote=quote))
        self._log(f"Received {len(matches)} matches.")
        return matches

    def _apply_highlights(self, document: Document, matches: list[QuoteMatch]) -> int:
        applied = 0
        for match in matches:
            color = CATEGORY_DETAILS[match.category]["color"]
            if highlight_quote(document, match.quote, color):
                applied += 1
        return applied

    # -------------------------------------------------------------- Utils --
    def _log(self, message: str) -> None:
        self.log.insert(END, f"{message}\n")
        self.log.see(END)

    def _append_legend(self, document: Document) -> None:
        legend_title = document.add_paragraph()
        legend_title_run = legend_title.add_run("Coding legend")
        legend_title_run.bold = True

        for code, meta in CATEGORY_DETAILS.items():
            entry = document.add_paragraph()
            entry.add_run(f"{code}. {meta['title']} – ")
            swatch = entry.add_run("example")
            apply_shading(swatch, meta["color"])

    @staticmethod
    def _build_client() -> AzureOpenAI:
        missing = []
        cleaned: dict[str, str] = {}
        for var in ("AZURE_OPENAI_API_KEY", "AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_DEPLOYMENT"):
            raw = os.environ.get(var, "").strip()
            if not raw:
                missing.append(var)
            else:
                cleaned[var] = raw
        if missing:
            message = "Missing Azure OpenAI environment variables: " + ", ".join(missing)
            raise EnvironmentError(message)

        endpoint = cleaned["AZURE_OPENAI_ENDPOINT"].rstrip("/")
        if not endpoint.startswith("https://"):
            raise EnvironmentError(
                "AZURE_OPENAI_ENDPOINT must include the full https URL, e.g. "
                "https://my-resource.openai.azure.com"
            )

        return AzureOpenAI(
            api_key=cleaned["AZURE_OPENAI_API_KEY"],
            azure_endpoint=endpoint,
            api_version=os.environ.get("AZURE_OPENAI_API_VERSION", "2024-02-15-preview"),
        )


# Helper functions ----------------------------------------------------------


def chunk_text(paragraphs: list[str], max_chars: int = 3500) -> list[str]:
    """Yield transcript chunks without splitting paragraphs."""

    chunks: list[str] = []
    buffer: list[str] = []
    current = 0
    for para in paragraphs:
        if current + len(para) + 1 > max_chars and buffer:
            chunks.append("\n".join(buffer))
            buffer = []
            current = 0
        buffer.append(para)
        current += len(para) + 1
    if buffer:
        chunks.append("\n".join(buffer))
    return chunks


def apply_shading(run, color_hex: str) -> None:
    """Apply a pastel shading color to a run using a hex fill value."""

    rpr = run._element.get_or_add_rPr()
    # Remove any existing highlight elements to avoid layering colors.
    for child in list(rpr):
        if child.tag == qn("w:highlight"):
            rpr.remove(child)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    rpr.append(shd)


def highlight_quote(document: Document, quote: str, color_hex: str) -> bool:
    """Highlight every paragraph containing the quote. Returns True if applied."""

    normalized = quote.strip()
    if not normalized:
        return False

    normalized_lower = normalized.lower()
    applied = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if normalized_lower in text.lower():
            for run in paragraph.runs:
                apply_shading(run, color_hex)
            applied = True
    return applied


def main() -> None:
    root = Tk()
    TranscriptCoderApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
